# -*- coding: utf-8 -*-

import sys
from pathlib import Path
from datetime import datetime
from copy import deepcopy

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DEFAULT_TEMPLATE_NAME = "接口测试报告模板.docx"


def get_base_dir():
    """
    兼容普通 Python 运行和 PyInstaller 打包后的 exe 运行。

    普通运行：
        模板从当前 py 文件所在目录读取。

    PyInstaller -F 打包后：
        模板会被解压到 sys._MEIPASS。
    """
    if getattr(sys, "frozen", False):
        return Path(sys._MEIPASS)
    else:
        return Path(__file__).resolve().parent


def normalize_text(text):
    """
    去掉换行、空格、制表符，方便识别 Word 表格里的内容。

    例如：
    avg
    /ms

    会被处理成：
    avg/ms
    """
    if text is None:
        return ""

    return (
        str(text)
        .replace("\n", "")
        .replace("\r", "")
        .replace("\t", "")
        .replace(" ", "")
        .strip()
    )


def replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    替换段落中的文本。

    优先在原 run 内替换，这样可以保留原字体、字号、加粗等样式。
    """
    if not old_text:
        return

    replaced = False

    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            replaced = True

    if replaced:
        return

    # 兜底：如果 XX 被 Word 拆成多个 run，则合并后替换。
    # 这种情况下会继承第一个 run 的样式。
    full_text = "".join(run.text for run in paragraph.runs)

    if old_text not in full_text:
        return

    new_full_text = full_text.replace(old_text, new_text)

    if paragraph.runs:
        paragraph.runs[0].text = new_full_text

        for run in paragraph.runs[1:]:
            run.text = ""


def replace_xx_in_doc(doc, third_party_name):
    """
    替换整篇 Word 文档里的 XX。
    包括：
    - 普通段落
    - 表格
    - 页眉
    - 页脚

    替换后尽量保留原有字体和字号。
    """
    if not third_party_name:
        return

    old_text = "XX"
    new_text = str(third_party_name)

    # 普通段落
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, old_text, new_text)

    # 表格里的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, old_text, new_text)

    # 页眉页脚
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text_in_paragraph(paragraph, old_text, new_text)

        for paragraph in section.footer.paragraphs:
            replace_text_in_paragraph(paragraph, old_text, new_text)

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, old_text, new_text)

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, old_text, new_text)


def ensure_rfonts(run):
    """
    确保 run 存在 rFonts 节点。
    """
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))

    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    return r_fonts


def set_run_font(run, font_name, font_size):
    """
    设置 run 的中文/英文/ASCII 字体和字号。
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)

    r_fonts = ensure_rfonts(run)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def set_cell_text_with_font(cell, value, font_name, font_size):
    """
    设置单元格文本，并指定字体和字号。
    font_size 单位是 pt。
    """
    value = "" if value is None else str(value)

    cell.text = ""

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)

    set_run_font(run, font_name, font_size)


def set_response_cell_text(cell, value):
    """
    响应时间表格数据字体：
    等线 Light，10号。
    """
    set_cell_text_with_font(cell, value, "等线 Light", 10)


def set_url_cell_text(cell, value):
    """
    测试URL、生产URL字体：
    仿宋，11号。
    """
    set_cell_text_with_font(cell, value, "仿宋", 11)


def set_paragraph_text_with_same_style(paragraph, text):
    """
    设置段落文本，并尽量沿用原段落第一个 run 的样式。
    用于附录里的查得入参、响应样例。
    """
    text = "" if text is None else str(text)

    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = text

        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def format_int_ms(value):
    """
    响应时间类字段：统一转成整数毫秒。
    """
    if value is None or value == "":
        return ""

    try:
        return str(int(round(float(value))))
    except Exception:
        return str(value)


def format_error_rate(value):
    """
    异常率：保留两位小数，默认加 %。
    """
    if value is None or value == "":
        return ""

    value_str = str(value).strip()

    if value_str.endswith("%"):
        value_str = value_str[:-1]

    try:
        return f"{float(value_str):.2f}%"
    except Exception:
        return str(value)


def format_tps(value):
    """
    TPS：保留 1 位小数。
    """
    if value is None or value == "":
        return ""

    try:
        return f"{float(value):.1f}"
    except Exception:
        return str(value)


def is_response_time_header_row(row):
    """
    判断某一行是不是响应时间表头。

    兼容：
    测试计划 | 条数 | avg/ms | mid/ms | P90/ms | ... | tps/s | 异常率
    """
    row_texts = [normalize_text(cell.text) for cell in row.cells]
    row_joined = "|".join(row_texts)

    required_keywords = [
        "条数",
        "avg",
        "mid",
        "P90",
        "P95",
        "P99",
        "min",
        "max",
        "tps",
        "异常率"
    ]

    matched_count = 0

    for keyword in required_keywords:
        if keyword in row_joined:
            matched_count += 1

    return matched_count == len(required_keywords)


def find_response_time_table(doc):
    """
    查找响应时间表格。
    不写死第几个表，通过表头字段定位。
    """
    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            if is_response_time_header_row(row):
                return table, row_index

    return None, None


def build_header_index(table, header_row_index):
    """
    根据响应时间表格的表头行，建立字段名和列下标的对应关系。
    """
    header_map = {}

    header_row = table.rows[header_row_index]

    for index, cell in enumerate(header_row.cells):
        text = normalize_text(cell.text)

        if text == "测试计划" or text == "数据源":
            header_map["test_plan"] = index
        elif text == "条数":
            header_map["request_count"] = index
        elif text.startswith("avg"):
            header_map["avg_response"] = index
        elif text.startswith("mid"):
            header_map["mid_response"] = index
        elif text.startswith("P90"):
            header_map["p90_response"] = index
        elif text.startswith("P95"):
            header_map["p95_response"] = index
        elif text.startswith("P99"):
            header_map["p99_response"] = index
        elif text.startswith("min"):
            header_map["min_response"] = index
        elif text.startswith("max"):
            header_map["max_response"] = index
        elif text.startswith("tps"):
            header_map["tps"] = index
        elif text == "异常率":
            header_map["error_rate"] = index

    return header_map


def get_first_data_row_index(table, header_row_index):
    """
    数据行从表头下一行开始。
    """
    return header_row_index + 1


def copy_row_style(source_row, target_row):
    """
    尽量复制模板数据行的单元格属性，让新增行的边框/宽度更接近模板。
    """
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        source_tc_pr = source_cell._tc.tcPr

        if source_tc_pr is not None:
            target_tc_pr = target_cell._tc.get_or_add_tcPr()
            target_tc_pr.clear()

            for child in source_tc_pr:
                target_tc_pr.append(deepcopy(child))


def ensure_data_rows(table, header_row_index, needed_count):
    """
    确保响应时间表格里有足够的数据行。
    如果模板只有一行空白数据行，但有多个 JTL，就自动新增行。
    """
    first_data_row_index = get_first_data_row_index(table, header_row_index)

    existing_count = len(table.rows) - first_data_row_index

    if existing_count <= 0:
        table.add_row()
        existing_count = 1

    template_row = table.rows[first_data_row_index]

    while existing_count < needed_count:
        new_row = table.add_row()
        copy_row_style(template_row, new_row)
        existing_count += 1

    return first_data_row_index


def clear_response_data_row(row, header_map):
    """
    清空响应时间表格某一行中需要填充的列。
    """
    for key, cell_index in header_map.items():
        if cell_index < len(row.cells):
            set_response_cell_text(row.cells[cell_index], "")


def find_source_info_table(doc):
    """
    查找“新源信息”表格。
    通过“测试URL”和“生产URL”两个字段定位。
    """
    for table in doc.tables:
        table_text = ""

        for row in table.rows:
            for cell in row.cells:
                table_text += normalize_text(cell.text) + "|"

        if "测试URL" in table_text and "生产URL" in table_text:
            return table

    return None


def set_value_after_label(table, label_text, value):
    """
    在表格中找到某个 label，比如“测试URL”，
    然后把它右侧的单元格填成 value。
    """
    target_label = normalize_text(label_text)

    for row in table.rows:
        cells = row.cells

        for index, cell in enumerate(cells):
            current_text = normalize_text(cell.text)

            if current_text == target_label:
                if index + 1 < len(cells):
                    set_url_cell_text(cells[index + 1], value)
                    return True

    return False


def fill_source_url_fields(doc, test_url, prod_url):
    """
    填充“新源信息”里的测试URL和生产URL。
    字体：仿宋，11号。
    """
    source_table = find_source_info_table(doc)

    if source_table is None:
        return {
            "success": False,
            "message": "未找到包含 测试URL、生产URL 的新源信息表格"
        }

    test_ok = set_value_after_label(source_table, "测试URL", test_url)
    prod_ok = set_value_after_label(source_table, "生产URL", prod_url)

    if not test_ok or not prod_ok:
        return {
            "success": False,
            "message": "已找到新源信息表格，但测试URL或生产URL字段填充失败"
        }

    return {
        "success": True,
        "message": "测试URL和生产URL填充成功"
    }

def set_date_cell_text(cell, value):
    """
    日期单元格字体。
    这里先用仿宋 11号，和测试URL/生产URL保持一致。
    """
    set_cell_text_with_font(cell, value, "仿宋", 11)


def fill_report_date(doc, date_text=None):
    """
    自动填充报告日期。

    只填充“日期”下面那一格：
        日期
        20260603

    不再填充“日期”右侧单元格。
    """

    if date_text is None:
        date_text = datetime.now().strftime("%Y%m%d")

    target_label = "日期"

    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                current_text = normalize_text(cell.text)

                if current_text == target_label:
                    # 必须有下一行，才填充下面那一格
                    if row_index + 1 >= len(table.rows):
                        return {
                            "success": False,
                            "message": "已找到“日期”，但它下面没有可填充的单元格"
                        }

                    below_cell = table.rows[row_index + 1].cells[col_index]

                    set_date_cell_text(below_cell, date_text)

                    return {
                        "success": True,
                        "message": f"报告日期已填充到“日期”下方单元格：{date_text}"
                    }

    return {
        "success": False,
        "message": "未找到文本为“日期”的表格单元格，未填充报告日期"
    }

def build_response_values(result_item, index):
    """
    构造一行响应时间表格的数据。
    """
    return {
        "test_plan": result_item.get("test_plan", f"文件{index}"),
        "request_count": str(result_item.get("request_count", "")),
        "avg_response": format_int_ms(result_item.get("avg_response", "")),
        "mid_response": format_int_ms(result_item.get("mid_response", "")),
        "p90_response": format_int_ms(result_item.get("p90_response", "")),
        "p95_response": format_int_ms(result_item.get("p95_response", "")),
        "p99_response": format_int_ms(result_item.get("p99_response", "")),
        "min_response": format_int_ms(result_item.get("min_response", "")),
        "max_response": format_int_ms(result_item.get("max_response", "")),
        "tps": format_tps(result_item.get("tps", "")),
        "error_rate": format_error_rate(result_item.get("error_rate", ""))
    }


def fill_response_time_table(table, header_row_index, result_list):
    """
    填充响应时间表格。
    多个 JTL 文件时，每个文件一行。
    第一列由 main.py 传入：
    - 第一个一般是“正式”
    - 后续一般是“文件2”“文件3”
    字体：等线 Light，10号。
    """
    header_map = build_header_index(table, header_row_index)

    if not result_list:
        return {
            "success": False,
            "message": "没有可填充的解析结果"
        }

    first_data_row_index = ensure_data_rows(
        table=table,
        header_row_index=header_row_index,
        needed_count=len(result_list)
    )

    for index, result_item in enumerate(result_list, start=1):
        row_index = first_data_row_index + index - 1
        row = table.rows[row_index]

        values = build_response_values(result_item, index)

        for key, value in values.items():
            if key not in header_map:
                continue

            cell_index = header_map[key]

            if cell_index < len(row.cells):
                set_response_cell_text(row.cells[cell_index], value)

    # 如果模板里原本有多余空白行，清空多余行，避免残留旧数据
    for row_index in range(first_data_row_index + len(result_list), len(table.rows)):
        clear_response_data_row(table.rows[row_index], header_map)

    return {
        "success": True,
        "message": "响应时间表格填充成功"
    }


def build_url_text(result_list):
    """
    多个 JTL 文件时，测试URL / 生产URL 只填一个 URL。
    默认取第一个非空 target_url。
    """
    url_list = []

    for item in result_list:
        target_url = item.get("target_url", "")

        if target_url and target_url not in url_list:
            url_list.append(target_url)

    if not url_list:
        return ""

    return url_list[0]


def get_first_sample_body(result_list):
    """
    从多个解析结果里取第一组非空 request_body / response_body。
    """
    for item in result_list:
        request_body = item.get("sample_request_body", "")
        response_body = item.get("sample_response_body", "")

        if request_body or response_body:
            return request_body, response_body

    return "", ""


def fill_appendix_sample_in_paragraphs(paragraphs, request_body, response_body):
    """
    在一组段落里填充：
    查得入参：
    响应样例：

    只填第一个“响应样例”，也就是“查得入参”下面的响应样例。
    """
    filled_request = False
    filled_response = False
    waiting_for_found_response = False

    for paragraph in paragraphs:
        text = normalize_text(paragraph.text)

        if not filled_request and text.startswith("查得入参"):
            new_text = "查得入参："

            if request_body:
                new_text += "\n" + request_body

            set_paragraph_text_with_same_style(paragraph, new_text)
            filled_request = True
            waiting_for_found_response = True
            continue

        if waiting_for_found_response and not filled_response and text.startswith("响应样例"):
            new_text = "响应样例："

            if response_body:
                new_text += "\n" + response_body

            set_paragraph_text_with_same_style(paragraph, new_text)
            filled_response = True
            waiting_for_found_response = False
            continue

        if text.startswith("查无入参"):
            waiting_for_found_response = False

    return filled_request, filled_response


def fill_appendix_sample(doc, request_body, response_body):
    """
    填充附录里的：
    查得入参：
    响应样例：

    模板里有两个“响应样例”，第一个属于“查得入参”，第二个属于“查无入参”。
    这里只填第一个响应样例。
    """
    if not request_body and not response_body:
        return

    filled_request, filled_response = fill_appendix_sample_in_paragraphs(
        doc.paragraphs,
        request_body,
        response_body
    )

    if filled_request and filled_response:
        return

    # 如果附录内容在表格里，再处理表格里的段落
    for table in doc.tables:
        table_paragraphs = []

        for row in table.rows:
            for cell in row.cells:
                table_paragraphs.extend(cell.paragraphs)

        sub_request, sub_response = fill_appendix_sample_in_paragraphs(
            table_paragraphs,
            request_body if not filled_request else "",
            response_body if not filled_response else ""
        )

        filled_request = filled_request or sub_request
        filled_response = filled_response or sub_response

        if filled_request and filled_response:
            return


def generate_report_from_template(result_data, template_path=None, output_path=None, third_party_name=""):
    """
    根据模板生成测试报告。

    result_data 可以是：
    1. 单个 dict
    2. 多个 dict 组成的 list

    多个 dict 时，响应时间表格每个 JTL 文件填一行。

    third_party_name：
    用户输入的三方名称，会替换模板中的所有 XX。
    """

    if isinstance(result_data, dict):
        result_list = [result_data]
    elif isinstance(result_data, list):
        result_list = result_data
    else:
        return {
            "success": False,
            "message": "result_data 类型错误，必须是 dict 或 list[dict]",
            "file_path": ""
        }

    if template_path is None:
        template_path = get_base_dir() / DEFAULT_TEMPLATE_NAME
    else:
        template_path = Path(template_path)

    if not template_path.exists():
        return {
            "success": False,
            "message": f"未找到模板文件：{template_path}",
            "file_path": ""
        }

    if output_path is None:
        output_dir = Path.cwd() / "output"
        output_dir.mkdir(exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = output_dir / f"接口测试报告_{timestamp}.docx"
    else:
        output_path = Path(output_path)

    doc = Document(str(template_path))

    # =========================
    # 0. 替换模板里的所有 XX
    # =========================
    replace_xx_in_doc(doc, third_party_name)

    # =========================
    # 0.1 填充报告日期
    # =========================
    date_fill_result = fill_report_date(doc)

    # =========================
    # 1. 填充响应时间表格
    # =========================
    table, header_row_index = find_response_time_table(doc)

    if table is None:
        return {
            "success": False,
            "message": "未找到响应时间表格，请确认模板中存在 条数、avg/ms、mid/ms、P90/ms、P95/ms、P99/ms、min/ms、max/ms、tps/s、异常率 这些表头",
            "file_path": ""
        }

    response_fill_result = fill_response_time_table(
        table=table,
        header_row_index=header_row_index,
        result_list=result_list
    )

    if not response_fill_result["success"]:
        return {
            "success": False,
            "message": response_fill_result["message"],
            "file_path": ""
        }

    # =========================
    # 2. 填充新源信息里的测试URL、生产URL
    # =========================
    url_text = build_url_text(result_list)

    if url_text:
        url_fill_result = fill_source_url_fields(
            doc=doc,
            test_url=url_text,
            prod_url=url_text
        )

        if not url_fill_result["success"]:
            return {
                "success": False,
                "message": url_fill_result["message"],
                "file_path": ""
            }

    # =========================
    # 3. 填充附录里的查得入参和响应样例
    # =========================
    sample_request_body, sample_response_body = get_first_sample_body(result_list)

    fill_appendix_sample(
        doc=doc,
        request_body=sample_request_body,
        response_body=sample_response_body
    )

    doc.save(str(output_path))

    return {
        "success": True,
        "message": "测试报告生成成功",
        "file_path": str(output_path),
        "file_name": output_path.name
    }